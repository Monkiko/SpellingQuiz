"""
The purpose of this program is to pull an indicated number of random words from a specified spelling list and write them to a new file while removing them from the original list.


Created by: Ian Rivera-Leandry
Last Updated: July 4, 2025
Version 2.2.6
"""



import os
from time import sleep
from random import shuffle
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor


# Get input from the user for the grade level of the desired spelling list and pass result to readSpellingList function
def start():
    grade_level = str(input("Please indicate the grade level spelling list you want (i.e. K, 1, 2, 3, etc.):  "))

    if len(grade_level) == 1:
        readSpellingList(grade_level)

    else:
        clean()
        print("Incorrect Input. Please enter either K, 1, 2, 3, etc.")
        sleep(3)
        start()


# Clear the screen dependent on host OS system
def clean():
    if os.name == "nt":
        _ = os.system("cls")
    else:
        _ = os.system("clear")


# Read the spelling list file for the grade level indicated previous and add the contents to a list for output printing
def readSpellingList(grade_level):

    match grade_level.lower():
        case "k":
            grade = "Kindergarten"
        case "1":
            grade = "First_Grade"
        case "2":
            grade = "Second_Grade"
        case "3":
            grade = "Third_Grade"
        case "4":
            grade = "Fourth_Grade"
        case "5":
            grade = "Fifth_Grade"
        case "6":
            grade = "Sixth_Grade"
        case "7":
            grade = "Seventh_Grade"
        case "8":
            grade = "Eighth_Grade"
    
    with open("Spelling_Lists/" + grade + "_List.txt", "r") as spelling_file:
        spellingList = []
        spellingList = spelling_file.readlines()

    spelling_file.close()

    shuffle(spellingList)
    printSimplifiedList(spellingList, grade)


# Creates a text file with a simplified list to be used for making spelling tests in the future
def printSimplifiedList(spellingList, grade):
    quiz_number = str(input("Please enter the quiz number for this spelling list: "))
    
    with open("Tests/" + grade + "/Simplified_Lists/Spelling_List_#" + quiz_number + ".txt", "w") as f:
        for i in range(10):
            f.write(spellingList[i])
    f.close()

    printSpellingList(spellingList, grade, quiz_number)


# Creates a new Word document and writes the spelling list to it in a numbered format
def printSpellingList(spellingList, grade, quiz_number):
    doc = Document()

    start_date = str(input("Please enter the start date for this spelling list (i.e. June 2): "))
    end_date = str(input("Please enter the end date for this spelling list (i.e. June 6): "))

    title = doc.add_heading("Spelling List #" + quiz_number, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.name = "Arial"
    title.style.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    title.style.font.size = Pt(26)

    date = doc.add_heading(start_date + " - " + end_date, level=2)
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date.style.font.name = "Arial"
    date.style.font.color.rgb = RGBColor(102, 102, 102)  # Dark gray color
    date.style.font.size = Pt(12)

    for i in range(10):
        list = doc.add_paragraph(str(i + 1) + ") " + spellingList[i].strip())
        list.style.font.name = "Arial"
        list.style.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        list.style.font.size = Pt(20)

    doc.save("Quizzes/" + grade + "/Spelling List #" + quiz_number + " - " + grade + ".docx")
    print("Simplified list saved to Tests/" + grade + "/Simplified_Lists/Spelling_List_#" + quiz_number + ".txt")
    print("Spelling list saved to Quizzes/" + grade + "/Spelling List #" + quiz_number + " - " + grade + ".docx")

    listCleanUp(spellingList, grade)


# Removes the words in the spelling list from the source Spelling_Lists file so words are not reused in future iterations
def listCleanUp(spellingList, grade):

    cleanlist = spellingList[10:]
          
    with open("Spelling_Lists/" + grade + "_List.txt", "w") as f:
        for i in cleanlist:
            f.writelines([i])

    f.close()

start()
