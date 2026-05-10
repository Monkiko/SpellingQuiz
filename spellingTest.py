#!/usr/bin/python3
"""
The purpose of this program is to create a Spelling Test document by pulling from the Simplified_Lists created from the Quizzes.


Created by: Ian Rivera-Leandry
Created on: May 3, 2026
Last Updated: May 10, 2026
Version 1.0.0
"""

import os
from time import sleep
from itertools import islice
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor

# Get input from the user for the grade level of the desired spelling test and pass result to readSpellingQuizzes function
def start():
    grade_level = str(input("Please indicate the grade level of the spelling test (i.e. K, 1, 2, 3, etc.):  "))

    if len(grade_level) == 1:
        match grade_level.lower():
            case "k":
                grade = "Kindergarten"
            case "1":
                grade = "First_Grade"
            case "2":
                grade = "Second_Grade"
            case "3":
                grade = "Third_Grade"
            case "4":
                grade = "Fourth_Grade"
            case "5":
                grade = "Fifth_Grade"
            case "6":
                grade = "Sixth_Grade"
            case "7":
                grade = "Seventh_Grade"
            case "8":
                grade = "Eighth_Grade"
        readSimplifiedLists(grade)

    else:
        clean()
        print("Incorrect Input. Please enter either K, 1, 2, 3, etc.")
        sleep(3)
        start()


# Clear the screen dependent on host OS system
def clean():
    if os.name == "nt":
        _ = os.system("cls")
    else:
        _ = os.system("clear")


def readSimplifiedLists(grade):
    spellingList = []
    folder_path = f'Tests/{grade}/Simplified_Lists/'
    quizzes = []
    test_number = str(input("Please indicate the test number (i.e. 1, 2, 3, etc.):  "))

    for i in range(4):
        print("Which quiz #'s are being used for this test? (i.e. 1, 2, 3, etc.):  ")
        quiz_number = str(input("Choice " + str(i + 1) + ":  "))
        quizzes.append(quiz_number)

    for quiz in quizzes:
        with open(folder_path + f'Spelling_List_#{quiz}.txt', 'r') as file:
            if int(test_number) % 2 != 0:
                first_five = [line.strip() for line in islice(file, 5)]
                spellingList.extend(first_five)
            else:
                last_five = file.readlines()[-5:]
                spellingList.extend([line.strip() for line in last_five])
    file.close()
    printSpellingTest(spellingList, grade, test_number)

def printSpellingTest(spellingList, grade, test_number):
    folder_path = f'Tests/{grade}/'

    doc = Document()

    start_date = str(input("Please enter the start date for this spelling list (i.e. June 2): "))
    end_date = str(input("Please enter the end date for this spelling list (i.e. June 6): "))
    
    title = doc.add_heading("Spelling Test #" + test_number, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.name = "Arial"
    title.style.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    title.style.font.size = Pt(26)

    date = doc.add_heading(start_date + " - " + end_date, level=2)
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date.style.font.name = "Arial"
    date.style.font.color.rgb = RGBColor(102, 102, 102)  # Dark gray color
    date.style.font.size = Pt(12)

    for i in range(20):
        list = doc.add_paragraph(str(i + 1) + ") " + spellingList[i])
        list.style.font.name = "Liberation Serif"
        list.style.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        list.style.font.size = Pt(12)

    doc.save(folder_path + f'Spelling_Test__#{test_number} - {grade}.docx')
    print("Spelling test saved to " + folder_path + f'Spelling_Test__#{test_number} - {grade}.docx')

start()